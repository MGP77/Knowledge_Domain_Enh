#!/usr/bin/env python3
"""
Сервис для обработки файлов (PDF, DOCX, TXT)

Copyright (c) 2025. All rights reserved.
Author: M.P.
"""

import os
import logging
import mimetypes
from typing import Optional, Dict, Any
from pathlib import Path

# Импорты для обработки различных типов файлов
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Проверяем доступность python-magic (может отсутствовать в корпоративной среде)
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

from config import config

logger = logging.getLogger(__name__)

class FileProcessorService:
    """Сервис для обработки загруженных файлов"""
    
    def __init__(self):
        self.upload_folder = config.UPLOAD_FOLDER
        self._ensure_upload_folder()
    
    def _ensure_upload_folder(self):
        """Создание папки для загрузок если не существует"""
        os.makedirs(self.upload_folder, exist_ok=True)
    
    def get_file_type(self, file_path: str) -> str:
        """
        Определение типа файла (корпоративная версия без libmagic)
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Тип файла
        """
        try:
            # Приоритет 1: python-magic (если доступен И libmagic установлен)
            if MAGIC_AVAILABLE:
                try:
                    file_type = magic.from_file(file_path, mime=True)
                    return file_type
                except Exception as magic_error:
                    # libmagic системная библиотека может отсутствовать
                    logger.warning(f"⚠️ python-magic установлен, но libmagic недоступен: {magic_error}")
                    logger.info("🔄 Переключаемся на встроенный mimetypes")
            
            # Приоритет 2: встроенный модуль mimetypes
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                logger.debug(f"📋 Тип файла определён через mimetypes: {mime_type}")
                return mime_type
                
            # Приоритет 3: по расширению файла
            extension = Path(file_path).suffix.lower()
            mime_types = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.txt': 'text/plain',
                '.md': 'text/markdown',
                '.html': 'text/html',
                '.htm': 'text/html'
            }
            detected_type = mime_types.get(extension, 'application/octet-stream')
            logger.debug(f"📁 Тип файла определён по расширению: {detected_type}")
            return detected_type
            
        except Exception as e:
            logger.warning(f"❌ Ошибка определения типа файла {file_path}: {e}")
            # Окончательный fallback - по расширению
            extension = Path(file_path).suffix.lower()
            if extension == '.pdf':
                return 'application/pdf'
            elif extension in ['.docx', '.doc']:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif extension in ['.txt', '.md']:
                return 'text/plain'
            else:
                return 'application/octet-stream'
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """
        Валидация загруженного файла
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Dict с результатом валидации
        """
        try:
            if not os.path.exists(file_path):
                return {
                    'valid': False,
                    'error': 'Файл не найден'
                }
            
            # Проверяем размер файла
            file_size = os.path.getsize(file_path)
            if file_size > config.MAX_FILE_SIZE:
                return {
                    'valid': False,
                    'error': f'Файл слишком большой. Максимальный размер: {config.MAX_FILE_SIZE // (1024*1024)} MB'
                }
            
            # Проверяем расширение
            file_extension = Path(file_path).suffix.lower().lstrip('.')
            if file_extension not in config.ALLOWED_EXTENSIONS:
                return {
                    'valid': False,
                    'error': f'Неподдерживаемый тип файла. Разрешены: {", ".join(config.ALLOWED_EXTENSIONS)}'
                }
            
            # Проверяем MIME тип
            file_type = self.get_file_type(file_path)
            
            return {
                'valid': True,
                'file_size': file_size,
                'file_type': file_type,
                'extension': file_extension
            }
            
        except Exception as e:
            logger.error(f"Ошибка валидации файла: {e}")
            return {
                'valid': False,
                'error': f'Ошибка валидации: {str(e)}'
            }
    
    def extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        """
        Извлечение текста из PDF файла
        
        Args:
            file_path: Путь к PDF файлу
            
        Returns:
            Извлеченный текст или None
        """
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 не установлен")
            return None
        
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из PDF: {e}")
            return None
    
    def extract_text_from_docx(self, file_path: str) -> Optional[str]:
        """
        Извлечение текста из DOCX файла
        
        Args:
            file_path: Путь к DOCX файлу
            
        Returns:
            Извлеченный текст или None
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx не установлен")
            return None
        
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Ошибка извлечения текста из DOCX: {e}")
            return None
    
    def extract_text_from_txt(self, file_path: str) -> Optional[str]:
        """
        Чтение текстового файла
        
        Args:
            file_path: Путь к TXT файлу
            
        Returns:
            Содержимое файла или None
        """
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return content.strip() if content.strip() else None
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Не удалось определить кодировку файла: {file_path}")
            return None
            
        except Exception as e:
            logger.error(f"Ошибка чтения текстового файла: {e}")
            return None
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Обработка файла и извлечение текста
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Dict с результатом обработки
        """
        try:
            # Валидация файла
            validation = self.validate_file(file_path)
            if not validation['valid']:
                return {
                    'success': False,
                    'error': validation['error']
                }
            
            file_extension = validation['extension']
            file_name = Path(file_path).name
            
            # Извлечение текста в зависимости от типа файла
            text_content = None
            
            if file_extension == 'pdf':
                text_content = self.extract_text_from_pdf(file_path)
            elif file_extension in ['docx', 'doc']:
                text_content = self.extract_text_from_docx(file_path)
            elif file_extension == 'txt':
                text_content = self.extract_text_from_txt(file_path)
            
            if not text_content:
                return {
                    'success': False,
                    'error': 'Не удалось извлечь текст из файла или файл пуст'
                }
            
            # Подготавливаем метаданные
            metadata = {
                'source': 'file',  # Стандартизировано для совместимости со статистикой
                'filename': file_name,
                'file_type': file_extension,
                'file_size': validation['file_size'],
                'title': file_name,
                'processed_at': str(Path(file_path).stat().st_mtime)
            }
            
            return {
                'success': True,
                'content': text_content,
                'metadata': metadata,
                'filename': file_name
            }
            
        except Exception as e:
            logger.error(f"Ошибка обработки файла: {e}")
            return {
                'success': False,
                'error': f'Ошибка обработки файла: {str(e)}'
            }
    
    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """
        Сохранение загруженного файла
        
        Args:
            file_content: Содержимое файла
            filename: Имя файла
            
        Returns:
            Путь к сохраненному файлу
        """
        try:
            file_path = os.path.join(self.upload_folder, filename)
            
            # Если файл уже существует, добавляем индекс
            counter = 1
            original_name, extension = os.path.splitext(filename)
            
            while os.path.exists(file_path):
                new_filename = f"{original_name}_{counter}{extension}"
                file_path = os.path.join(self.upload_folder, new_filename)
                counter += 1
            
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"✅ Файл сохранен: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Ошибка сохранения файла: {e}")
            raise
